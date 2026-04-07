#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工时评估报告Word生成器

依赖安装：
    pip install python-docx

使用方式：
    python generate_word_report.py

或者：
    from generate_word_report import ReportGenerator
    generator = ReportGenerator()
    generator.generate(output_path="工时评估报告.docx")
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime


# 主色调
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)       # 深蓝 - 一级标题
COLOR_SECONDARY = RGBColor(0x2E, 0x75, 0xB6)      # 中蓝 - 二级标题
COLOR_TABLE_HEADER_HEX = '1F4E79'                  # 深蓝 - 表头背景(hex)
COLOR_TABLE_ALT = 'F2F7FB'                          # 浅蓝 - 表格斑马纹
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)            # 深灰 - 正文
COLOR_NOTE = RGBColor(0x66, 0x66, 0x66)            # 中灰 - 注释/页脚


class ReportGenerator:
    """工时评估报告生成器"""

    def __init__(self):
        self.doc = Document()

    def _set_style(self):
        """设置文档全局样式"""
        # 页边距
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        # 正文样式
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(11)
        style.font.color.rgb = COLOR_BODY
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # 标题样式
        for level, (font_size, color, font_name) in {
            1: (Pt(16), COLOR_PRIMARY, '微软雅黑'),
            2: (Pt(13), COLOR_SECONDARY, '微软雅黑'),
        }.items():
            h_style = self.doc.styles[f'Heading {level}']
            h_style.font.name = font_name
            h_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            h_style.font.size = font_size
            h_style.font.color.rgb = color
            h_style.font.bold = True
            h_style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            h_style.paragraph_format.space_after = Pt(8)

    def _add_heading(self, text, level=1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 一级标题下方加蓝色细线
        if level == 1:
            pPr = heading._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="1F4E79"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
        return heading

    def _add_paragraph(self, text, bold=False, size=11, color=None, align=None):
        """添加段落"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.color.rgb = color or COLOR_BODY
        if align:
            para.alignment = align
        return para

    def _add_table(self, headers, rows):
        """添加美观表格：深蓝表头白字 + 斑马纹"""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 表头：深蓝背景白字
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER_HEX}" w:val="clear"/>'
            )
            cell._element.get_or_add_tcPr().append(shading)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header)
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 数据行：斑马纹
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ''
                # 偶数行浅蓝背景
                if row_idx % 2 == 1:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_ALT}" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(str(cell_data))
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.color.rgb = COLOR_BODY

        # 设置单元格垂直居中和内边距
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
                cell.paragraphs[0].paragraph_format.space_before = Pt(3)
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)

        self.doc.add_paragraph()  # 表格后留白
        return table

    def _add_bullet_list(self, items, level=0):
        """添加项目符号列表"""
        for item in items:
            para = self.doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(item)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.color.rgb = COLOR_BODY

    def _add_cover(self, title):
        """添加封面页"""
        # 顶部留白
        for _ in range(6):
            self.doc.add_paragraph()

        # 标题
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = COLOR_PRIMARY

        # 蓝色分隔线
        line_para = self.doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = line_para._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="6" w:color="2E75B6"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        self.doc.add_paragraph()

        # 日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(datetime.now().strftime('%Y 年 %m 月 %d 日'))
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.color.rgb = COLOR_NOTE

        # 分页
        self.doc.add_page_break()

    def generate(self, output_path="工时评估报告.docx", data=None):
        """
        生成工时评估报告

        Args:
            output_path: 输出文件路径
            data: 报告数据（可选，默认使用内置示例数据）
        """
        self._set_style()

        if data is None:
            data = self._get_sample_data()

        # 封面
        self._add_cover(data.get('title', '售前工时评估报告'))

        # 一、需求概述
        self._add_heading('一、需求概述', level=1)

        self._add_heading('1.1 需求背景', level=2)
        self._add_paragraph(data.get('background', '（描述客户提出的原始需求背景）'))

        self._add_heading('1.2 核心痛点', level=2)
        self._add_bullet_list(data.get('pain_points', [
            '痛点1：xxx',
            '痛点2：xxx'
        ]))

        self._add_heading('1.3 潜在需求', level=2)
        self._add_bullet_list(data.get('latent_needs', [
            '潜在需求1：xxx',
            '潜在需求2：xxx'
        ]))

        self._add_heading('1.4 风险提示', level=2)
        self._add_bullet_list(data.get('risks', [
            '风险1：xxx（技术/业务/集成）',
            '风险2：xxx'
        ]))

        # 二、功能范围
        self._add_heading('二、功能范围', level=1)

        self._add_heading('2.1 用户故事拆解', level=2)
        self._add_table(
            ['用户故事', '优先级', '复杂度'],
            data.get('user_stories', [
                ['作为[角色]，我希望[行为]，从而[目的]', 'P1', '中']
            ])
        )

        self._add_heading('2.2 功能清单', level=2)
        self._add_bullet_list(data.get('feature_list', [
            '功能1：xxx',
            '功能2：xxx'
        ]))

        # 三、技术方案
        self._add_heading('三、技术方案', level=1)

        self._add_heading('3.1 涉及模块', level=2)
        self._add_bullet_list(data.get('modules', [
            '模块1：xxx',
            '模块2：xxx'
        ]))

        self._add_heading('3.2 技术栈', level=2)
        tech_stack = data.get('tech_stack', {
            '前端': 'Vue / JavaScript / CSS',
            '后端': 'Java / Spring Boot',
            '数据库': 'PostgreSQL / MySQL / Redis',
            '中间件': 'xxx'
        })
        for key, value in tech_stack.items():
            para = self._add_paragraph(f'{value}')
            run = para.runs[0]
            # 在前面插入加粗的 key
            para.clear()
            key_run = para.add_run(f'{key}：')
            key_run.bold = True
            key_run.font.size = Pt(11)
            key_run.font.name = '宋体'
            key_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            key_run.font.color.rgb = COLOR_BODY
            val_run = para.add_run(value)
            val_run.font.size = Pt(11)
            val_run.font.name = '宋体'
            val_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            val_run.font.color.rgb = COLOR_BODY

        self._add_heading('3.3 影响范围', level=2)
        impact_scope = data.get('impact_scope', {})
        for label, key in [('涉及对象/字段', 'objects'), ('涉及接口', 'interfaces'), ('涉及流程', 'processes')]:
            para = self.doc.add_paragraph()
            key_run = para.add_run(f'{label}：')
            key_run.bold = True
            key_run.font.size = Pt(11)
            key_run.font.name = '宋体'
            key_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            key_run.font.color.rgb = COLOR_BODY
            val_run = para.add_run(impact_scope.get(key, 'xxx'))
            val_run.font.size = Pt(11)
            val_run.font.name = '宋体'
            val_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            val_run.font.color.rgb = COLOR_BODY

        # 四、资源配置
        self._add_heading('四、资源配置', level=1)

        self._add_heading('4.1 产品设计', level=2)
        self._add_table(
            ['姓名', '技能等级', '负责内容', '预估投入'],
            data.get('designers', [])
        )

        self._add_heading('4.2 前端开发', level=2)
        self._add_table(
            ['姓名', '技能等级', '负责内容', '预估投入'],
            data.get('frontend_devs', [])
        )

        self._add_heading('4.3 后端开发', level=2)
        self._add_table(
            ['姓名', '技能等级', '负责内容', '预估投入'],
            data.get('backend_devs', [])
        )

        self._add_heading('4.4 测试人员', level=2)
        self._add_table(
            ['姓名', '负责模块', '测试类型', '预估投入'],
            data.get('testers', [])
        )

        self._add_heading('4.5 资源风险', level=2)
        self._add_bullet_list(data.get('resource_risks', [
            'xxx'
        ]))

        # 五、工时评估
        self._add_heading('五、工时评估', level=1)

        self._add_heading('5.1 工时明细', level=2)
        self._add_table(
            ['阶段', '任务', '设计', '前端', '后端', '测试', '合计(人天)'],
            data.get('work_hours', [])
        )

        self._add_heading('5.2 工时说明', level=2)
        self._add_bullet_list(data.get('work_notes', [
            '按技能等级3（精通）人员标准估算',
            '技能等级2人员需增加30%工时',
            '技能等级1人员需增加50%工时',
            '包含需求确认、开发、测试、上线全过程',
            '不包含需求变更和返工时间'
        ]))

        self._add_heading('5.3 代码复杂度评估', level=2)
        complexity = data.get('complexity', {})
        self._add_paragraph('现有可复用代码：' + complexity.get('reusable', 'xxx（节省约xx人天）'))
        self._add_paragraph('新增代码量预估：约 ' + complexity.get('new_code', 'xxx 行'))
        self._add_paragraph('复杂模块说明：' + complexity.get('complex_modules', 'xxx'))

        # 六、待确认项
        self._add_heading('六、待确认项', level=1)
        self._add_bullet_list(data.get('confirm_items', [
            '确认项1：xxx',
            '确认项2：xxx'
        ]))

        # 七、报价建议
        self._add_heading('七、报价建议', level=1)
        self._add_table(
            ['项目', '工时(人天)', '单价(元)', '小计(元)'],
            data.get('quotation', [])
        )
        self._add_paragraph(
            '注：具体报价需根据客户类型、项目规模、合作模式等确定。',
            size=10, color=COLOR_NOTE
        )

        # 页脚
        self.doc.add_paragraph()
        self._add_paragraph(
            f'报告生成时间：{datetime.now().strftime("%Y-%m-%d")}    评估工具：AI工时评估助手',
            size=9, color=COLOR_NOTE, align=WD_ALIGN_PARAGRAPH.CENTER
        )

        # 保存文档
        self.doc.save(output_path)
        print(f"报告已生成：{output_path}")
        return output_path

    def _get_sample_data(self):
        """获取示例数据"""
        return {
            'title': '数据迁移工具开发 - 工时评估报告',
            'background': '客户需要将旧系统中多个业务模块的历史数据，通过自动化方式迁移到新的管理系统中。',
            'pain_points': [
                '客户有大量历史数据存在于旧系统中',
                '手工迁移成本高、风险大、容易出错',
                '需要确保数据完整性、一致性和可追溯性'
            ],
            'latent_needs': [
                '数据迁移工具需具备可复用性，以支持未来其他系统的数据迁移',
                '需要数据映射配置的可视化管理界面',
                '迁移过程需要监控和日志记录',
                '迁移失败后的回滚机制'
            ],
            'risks': [
                '技术风险：源系统表结构未知，存在数据格式差异风险',
                '数据风险：历史数据质量未知，可能存在脏数据、不一致问题',
                '业务风险：迁移期间影响系统可用性',
                '集成风险：与源系统的对接方式需确认（数据库直连/API接口）'
            ],
            'user_stories': [
                ['作为系统管理员，我希望配置源系统到目标系统的数据映射规则，从而灵活控制数据转换逻辑', 'P1', '高'],
                ['作为系统管理员，我希望一键启动数据迁移任务，从而高效完成历史数据迁移', 'P1', '中'],
                ['作为系统管理员，我希望迁移过程支持断点续传，从而处理网络中断或系统故障的情况', 'P1', '中'],
                ['作为系统管理员，我希望迁移完成后自动生成对账清单，从而验证数据一致性', 'P1', '中'],
                ['作为系统管理员，我希望迁移包含附件文档，从而保持完整的历史记录', 'P1', '中']
            ],
            'feature_list': [
                '数据映射配置功能（表级映射、字段级映射、转换规则配置）',
                '一键式数据迁移功能（支持多个业务模块）',
                '断点续传功能（记录迁移进度、支持中断后继续）',
                '数据校验功能（MD5校验、记录行数比对、对账清单生成）',
                '附件迁移功能（业务单据、审核日志、电子签名及关联附件）',
                '迁移日志与进度展示',
                '回滚机制'
            ],
            'modules': [
                '业务模块A：核心业务数据',
                '业务模块B：项目管理',
                '系统集成：外部数据对接'
            ],
            'tech_stack': {
                '前端': 'Vue 2.x / JavaScript / CSS',
                '后端': 'Java 8+ / Spring Boot / MyBatis-Plus',
                '数据库': 'PostgreSQL / MySQL / Redis（缓存迁移状态）',
                '中间件': 'Minio（附件存储）、Quartz（定时任务）'
            },
            'impact_scope': {
                'objects': '业务实例、项目实例、关联附件、审核日志、电子签名',
                'interfaces': '新增独立的数据迁移工具接口（不涉及通用CRUD）',
                'processes': '数据映射配置、迁移任务执行、断点续传、对账清单生成'
            },
            'frontend_devs': [
                ['张三', '3', '迁移配置界面、进度展示页、对账清单导出页', '15 人天'],
                ['李四', '2', '迁移任务管理、日志展示', '10 人天']
            ],
            'backend_devs': [
                ['王五', '3', '迁移工具核心引擎、断点续传、MD5校验', '25 人天'],
                ['赵六', '3', '迁移方案设计、数据库直连对接、附件迁移', '10 人天'],
                ['孙七', '2', '数据映射规则、业务逻辑适配', '15 人天（×1.3）'],
                ['周八', '2', '项目数据映射规则、业务逻辑适配', '8 人天（×1.3）']
            ],
            'testers': [
                ['吴九', '通用功能', '功能测试、接口自动化测试', '15 人天'],
                ['郑十', '项目管理', '功能测试、集成测试', '10 人天']
            ],
            'designers': [
                ['陈设计', '3', '需求澄清、原型设计、交互评审', '5 人天']
            ],
            'resource_risks': [
                '源系统对接方式未确定，需赵六评估可行性',
                '孙七、周八技能等级为2（熟悉），实际工时需增加30%',
                '附件迁移涉及Minio存储，需确认源系统附件存储方式',
                '数据量预估未提供，大规模数据迁移可能需要性能优化'
            ],
            'work_hours': [
                ['需求分析', '需求澄清、源系统调研、方案设计', '5', '2', '5', '-', '12'],
                ['前端开发', '配置界面、进度展示、对账清单', '-', '25', '-', '-', '25'],
                ['后端开发', '核心引擎、映射规则、迁移逻辑', '-', '-', '58', '-', '58'],
                ['数据库', '迁移配置表、迁移日志表', '-', '-', '2', '-', '2'],
                ['系统集成', '源系统对接、附件迁移', '-', '-', '8', '-', '8'],
                ['测试', '功能测试、集成测试、回归测试', '-', '-', '-', '25', '25'],
                ['部署', '环境配置、上线验证', '-', '-', '3', '-', '3'],
                ['合计', '', '5', '27', '76', '25', '133']
            ],
            'work_notes': [
                '按技能等级3（精通）人员标准估算',
                '技能等级2人员需增加30%工时（孙七、周八已调整）',
                '需求分析阶段工时归入产品设计（"设计"列）',
                '包含需求确认、开发、测试、上线全过程',
                '不包含需求变更和返工时间',
                '不包含源系统数据清洗和预处理时间'
            ],
            'complexity': {
                'reusable': '文件处理工具类、Excel导入导出基础服务（节省约10人天）',
                'new_code': '5,000-6,000',
                'complex_modules': '断点续传机制、MD5校验（大数据量需考虑性能优化）、附件迁移（跨系统，需处理加密/解密、格式转换）'
            },
            'confirm_items': [
                '源系统是否提供数据库直连权限？还是通过API对接？',
                '源系统各模块的具体表结构和字段清单需客户提供',
                '迁移期间是否允许用户继续使用系统？是否需要停机维护窗口？',
                '迁移失败后的回滚范围：全量回滚还是单表回滚？',
                '数据量预估（各模块记录数、附件大小），用于评估迁移时长'
            ],
            'quotation': [
                ['设计费用', '5', '2,500', '12,500'],
                ['开发费用', '103', '3,000', '309,000'],
                ['测试费用', '25', '2,000', '50,000'],
                ['实施费用', '10', '3,000', '30,000'],
                ['培训费用', '3', '2,000', '6,000'],
                ['合计', '146', '', '407,500']
            ]
        }


def generate_from_template(output_path="工时评估报告.docx", data=None):
    """
    从模板数据生成报告的便捷函数

    Args:
        output_path: 输出文件路径
        data: 自定义报告数据，如为None则使用示例数据
    """
    generator = ReportGenerator()
    return generator.generate(output_path, data)


if __name__ == '__main__':
    # 生成示例报告
    generate_from_template()
